from io import BytesIO

from jobhunt.profile.cv_extract import extract_text


def test_extract_text_from_txt():
    out = extract_text("cv.txt", b"Python Backend Engineer")
    assert "Python Backend Engineer" in out


def test_extract_text_from_docx():
    import docx

    document = docx.Document()
    document.add_paragraph("Python Backend Engineer")
    document.add_paragraph("Skills: FastAPI, AWS")
    buffer = BytesIO()
    document.save(buffer)

    out = extract_text("cv.docx", buffer.getvalue())
    assert "Python Backend Engineer" in out
    assert "FastAPI" in out
